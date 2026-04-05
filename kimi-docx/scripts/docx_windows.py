from __future__ import annotations

import json
import sys
from pathlib import Path

import docx
import lxml
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


def set_run_font(run, east_asia: str = "Microsoft YaHei", ascii_font: str = "Calibri", size: int = 11, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def env() -> int:
    print(json.dumps({"python": sys.executable, "python_docx_module": getattr(docx, "__file__", None), "lxml_version": getattr(lxml, "__version__", None)}, ensure_ascii=False, indent=2))
    return 0


def _add_paragraph(document: Document, text: str, style: str | None = None) -> None:
    paragraph = document.add_paragraph()
    if style:
        try:
            paragraph.style = style
        except KeyError:
            pass
    set_run_font(paragraph.add_run(text))


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    try:
        paragraph.style = f"Heading {level}"
    except KeyError:
        pass
    set_run_font(paragraph.add_run(text), size=max(12, 18 - level * 2), bold=True)


def _add_table(document: Document, table_data: dict) -> None:
    headers = table_data.get("headers", [])
    rows = table_data.get("rows", [])
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers):
        set_run_font(cell.paragraphs[0].add_run(str(text)), bold=True)
    for row_values in rows:
        row = table.add_row().cells
        for cell, text in zip(row, row_values):
            set_run_font(cell.paragraphs[0].add_run(str(text)))


def read_summary(path: Path) -> int:
    document = Document(path)
    headings = []
    for paragraph in document.paragraphs:
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            headings.append(paragraph.text)
    print(json.dumps({"paragraphs": len(document.paragraphs), "tables": len(document.tables), "headings": headings}, ensure_ascii=False, indent=2))
    return 0


def create_report(input_path: Path, output_path: Path) -> int:
    payload = json.loads(input_path.read_text(encoding="utf-8"))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    set_run_font(document.sections[0].header.paragraphs[0].add_run(payload.get("header", "Kimi WSL Docx")), size=9)
    title = document.add_paragraph()
    set_run_font(title.add_run(payload.get("title", "Untitled")), size=18, bold=True)
    for item in payload.get("paragraphs", []):
        _add_paragraph(document, str(item))
    for table_data in payload.get("tables", []):
        _add_table(document, table_data)
    image_path = payload.get("image_path")
    if image_path:
        document.add_picture(str(image_path), width=Inches(1.0))
    document.save(output_path)
    print(output_path)
    return 0


def modify_brief(input_path: Path, output_path: Path) -> int:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document(input_path)
    _add_heading(document, "Updated Section", level=1)
    _add_paragraph(document, "Revision: Added launch readiness notes and refreshed ownership.")
    if document.tables and len(document.tables[0].rows) > 1 and len(document.tables[0].rows[1].cells) > 1:
        document.tables[0].rows[1].cells[1].text = "Ravi (Updated)"
    document.save(output_path)
    print(output_path)
    return 0


def create_advanced(input_json: Path, image_path: Path, output_path: Path) -> int:
    payload = json.loads(input_json.read_text(encoding="utf-8"))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    set_run_font(section.header.paragraphs[0].add_run(f"{payload['project_name']} Header"), size=9)
    set_run_font(section.footer.paragraphs[0].add_run("Confidential"), size=9)
    _add_heading(document, f"{payload['project_name']} Report", level=0)
    _add_heading(document, "Objectives", level=1)
    for item in payload.get("objectives", []):
        paragraph = document.add_paragraph(style="List Number")
        set_run_font(paragraph.add_run(item))
    _add_heading(document, "Actions", level=1)
    _add_table(document, {"headers": ["Item", "Owner", "Status"], "rows": [[a["item"], a["owner"], a["status"]] for a in payload.get("actions", [])]})
    document.add_picture(str(image_path), width=Inches(1.0))
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    _add_heading(document, "Risks", level=1)
    for risk in payload.get("risks", []):
        _add_paragraph(document, risk)
    document.save(output_path)
    print(output_path)
    return 0


def export_pdf(input_path: Path, output_path: Path) -> int:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document(input_path)
    styles = getSampleStyleSheet()
    story = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            style = styles["Heading2"] if paragraph.style and paragraph.style.name.startswith("Heading") else styles["BodyText"]
            story.append(Paragraph(text, style))
            story.append(Spacer(1, 8))
    for table_obj in document.tables:
        rows = [[cell.text for cell in row.cells] for row in table_obj.rows]
        table = Table(rows)
        table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("GRID", (0, 0), (-1, -1), 0.5, colors.grey)]))
        story.append(table)
        story.append(Spacer(1, 12))
    SimpleDocTemplate(str(output_path), pagesize=A4).build(story)
    print(output_path)
    return 0


def validate(path: Path) -> int:
    document = Document(path)
    print(json.dumps({"paragraphs": len(document.paragraphs), "tables": len(document.tables), "sections": len(document.sections)}, ensure_ascii=False, indent=2))
    return 0


def main() -> int:
    if len(sys.argv) < 2:
        return 1
    command = sys.argv[1]
    if command == "env":
        return env()
    if command == "read-summary":
        return read_summary(Path(sys.argv[2]))
    if command == "create-report":
        return create_report(Path(sys.argv[2]), Path(sys.argv[3]))
    if command == "modify-brief":
        return modify_brief(Path(sys.argv[2]), Path(sys.argv[3]))
    if command == "create-advanced":
        return create_advanced(Path(sys.argv[2]), Path(sys.argv[3]), Path(sys.argv[4]))
    if command == "export-pdf":
        return export_pdf(Path(sys.argv[2]), Path(sys.argv[3]))
    if command == "validate":
        return validate(Path(sys.argv[2]))
    return 1


if __name__ == "__main__":
    raise SystemExit(main())
